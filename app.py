import streamlit as st
import groq
import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import json
import re
from datetime import datetime

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
GROQ_API_KEY    = st.secrets.get("GROQ_API_KEY", "")
RAZORPAY_BASIC  = st.secrets.get("RAZORPAY_BASIC", "#")
RAZORPAY_PRO    = st.secrets.get("RAZORPAY_PRO", "#")
RAZORPAY_YEARLY = st.secrets.get("RAZORPAY_YEARLY", "#")
EMAIL_DB_FILE   = "used_emails.json"

SHEET_SCRIPT_URL = "https://script.google.com/macros/s/AKfycbzN99UwHn4Bt1mJ4MMS5ZSV-cysoTC_ac6d6oMNkWB_JAGb1i2vqBX3RmrCDqIsla3G/exec"

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="ResumeReflect — AI Resume Tailor",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ─────────────────────────────────────────────
# PREMIUM DARK UI
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,400&display=swap');

:root {
    --bg:      #fffdf5;
    --surface: #fff8e1;
    --card:    #ffffff;
    --border:  #f0d060;
    --accent:  #c9960c;
    --accent2: #f5c518;
    --green:   #b8860b;
    --text:    #1a1200;
    --muted:   #8a7030;
    --radius:  14px;
}

html, body, .stApp { background: var(--bg) !important; color: var(--text) !important; font-family: 'DM Sans', sans-serif !important; }
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 1.5rem 2.5rem !important; max-width: 1000px !important; margin: 0 auto; }
h1,h2,h3 { font-family: 'Syne', sans-serif !important; }

.hero { text-align: center; padding: 3rem 1rem 1.5rem; }
.badge {
    display: inline-block;
    background: linear-gradient(135deg,#f5c51822,#c9960c22);
    border: 1px solid #c9960c44;
    border-radius: 100px; padding: 5px 16px;
    font-size: 11px; letter-spacing: 2px; text-transform: uppercase;
    color: var(--accent); margin-bottom: 1.2rem;
}
.hero h1 {
    font-size: clamp(2.2rem,5vw,3.5rem) !important;
    font-weight: 800 !important; line-height: 1.1 !important; margin: 0 0 .8rem !important;
    background: linear-gradient(135deg,#1a1200 30%,#c9960c);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
}
.hero p { color: var(--muted); font-size: 1rem; max-width: 480px; margin: 0 auto 2rem; line-height: 1.7; }

.card {
    background: var(--card); border: 1px solid var(--border);
    border-radius: var(--radius); padding: 1.5rem; margin-bottom: 1rem;
    transition: border-color .2s; box-shadow: 0 2px 12px #f5c51820;
}
.card:hover { border-color: #c9960c88; }
.card-title {
    font-family: 'Syne', sans-serif; font-weight: 700; font-size: .95rem;
    color: var(--text); margin-bottom: 1rem; display: flex; align-items: center; gap: 8px;
}
.step-num {
    background: linear-gradient(135deg,var(--accent),var(--accent2));
    color: white; border-radius: 50%; width: 26px; height: 26px;
    display: inline-flex; align-items: center; justify-content: center;
    font-size: 12px; font-weight: 700; flex-shrink: 0;
}

.stFileUploader > div {
    background: var(--surface) !important; border: 2px dashed var(--border) !important;
    border-radius: 12px !important; color: var(--muted) !important;
}
.stFileUploader > div:hover { border-color: var(--accent) !important; }
.stFileUploader button {
    background: linear-gradient(135deg,var(--accent),var(--accent2)) !important;
    color: white !important; border: none !important;
    border-radius: 8px !important; font-weight: 700 !important;
}
.stFileUploader label { color: var(--text) !important; }
.stTextArea textarea {
    background: var(--surface) !important; border: 1px solid var(--border) !important;
    border-radius: 12px !important; color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important; font-size: 14px !important;
}
.stTextArea textarea:focus { border-color: var(--accent) !important; box-shadow: 0 0 0 2px #c9960c22 !important; }
.stTextInput input {
    background: var(--surface) !important; border: 1px solid var(--border) !important;
    border-radius: 10px !important; color: var(--text) !important;
}
.stTextInput input:focus { border-color: var(--accent) !important; }

.stButton > button {
    background: linear-gradient(135deg,var(--accent),var(--accent2)) !important;
    color: white !important; border: none !important; border-radius: 12px !important;
    padding: 13px 28px !important; font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; font-size: 15px !important; width: 100% !important;
    transition: all .2s !important;
}
.stButton > button:hover { transform: translateY(-2px) !important; box-shadow: 0 8px 25px #c9960c44 !important; }
.stDownloadButton > button {
    background: linear-gradient(135deg,var(--accent),var(--accent2)) !important;
    color: white !important; border: none !important; border-radius: 12px !important;
    padding: 13px 28px !important; font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; font-size: 15px !important; width: 100% !important;
    transition: all .2s !important;
}
.stDownloadButton > button:hover { transform: translateY(-2px) !important; box-shadow: 0 8px 25px #c9960c44 !important; }

.pricing-grid { display: grid; grid-template-columns: repeat(3,1fr); gap: 1rem; margin: 1.2rem 0; }
.plan {
    background: var(--card); border: 1px solid var(--border);
    border-radius: var(--radius); padding: 1.4rem; text-align: center;
    transition: all .2s; position: relative; box-shadow: 0 2px 12px #f5c51820;
}
.plan:hover { border-color: var(--accent); transform: translateY(-3px); box-shadow: 0 10px 28px #c9960c22; }
.plan.hot { border-color: var(--accent); background: linear-gradient(135deg,#f5c51810,var(--card)); }
.hot-badge {
    position: absolute; top: -11px; left: 50%; transform: translateX(-50%);
    background: linear-gradient(135deg,var(--accent),var(--accent2));
    color: white; padding: 2px 12px; border-radius: 100px;
    font-size: 10px; font-weight: 700; letter-spacing: 1px; text-transform: uppercase;
}
.plan-name { font-family:'Syne',sans-serif; font-weight:700; font-size:.85rem; color:var(--muted); text-transform:uppercase; letter-spacing:1.5px; margin-bottom:.4rem; }
.plan-price { font-family:'Syne',sans-serif; font-weight:800; font-size:1.8rem; color:var(--accent); }
.plan-period { font-size:11px; color:var(--muted); margin-bottom:.8rem; }
.plan-feat { list-style:none; padding:0; margin:0 0 1rem; text-align:left; }
.plan-feat li { font-size:12px; color:var(--muted); padding:3px 0; display:flex; align-items:center; gap:6px; }
.plan-feat li::before { content:"✓"; color:var(--accent); font-weight:700; }
.pay-btn {
    display:block; background:linear-gradient(135deg,var(--accent),var(--accent2));
    color:white !important; text-decoration:none !important; border-radius:9px;
    padding:9px 16px; font-family:'Syne',sans-serif; font-weight:700;
    font-size:13px; transition:all .2s; text-align:center;
}
.pay-btn:hover { box-shadow:0 5px 18px #c9960c44; transform:translateY(-1px); }

.score-wrap { display:flex; gap:1rem; margin:1rem 0; }
.score-box {
    flex:1; background:var(--card); border:1px solid var(--border);
    border-radius:var(--radius); padding:1.2rem; text-align:center;
    box-shadow: 0 2px 12px #f5c51820;
}
.score-val { font-family:'Syne',sans-serif; font-weight:800; font-size:2.2rem; }
.score-before { background:linear-gradient(135deg,#8a7030,#c9a84c); -webkit-background-clip:text; -webkit-text-fill-color:transparent; background-clip:text; }
.score-after  { background:linear-gradient(135deg,var(--accent),var(--accent2)); -webkit-background-clip:text; -webkit-text-fill-color:transparent; background-clip:text; }
.score-label  { font-size:11px; color:var(--muted); margin-top:.3rem; text-transform:uppercase; letter-spacing:1px; }
.bar-wrap { background:var(--surface); border-radius:100px; height:8px; margin:.4rem 0; overflow:hidden; }
.bar { height:100%; border-radius:100px; background:linear-gradient(90deg,var(--accent),var(--accent2)); }

.chip { display:inline-block; background:#f5c51822; border:1px solid #c9960c33; color:var(--accent); border-radius:100px; padding:3px 10px; font-size:11px; margin:2px; }
.divider { height:1px; background:linear-gradient(90deg,transparent,var(--border),transparent); margin:2rem 0; }
.unlock-box { background: var(--card); border: 1px solid #c9960c33; border-radius: var(--radius); padding: 1.5rem; margin-top: 1rem; }

.stRadio > div { gap: 10px !important; }
.stRadio label { background:var(--surface) !important; border:1px solid var(--border) !important; border-radius:9px !important; padding:9px 14px !important; cursor:pointer !important; transition:all .15s !important; }
.stRadio label:hover { border-color:var(--accent) !important; }
label, .stRadio p { color:var(--muted) !important; font-size:13px !important; }
.stSelectbox > div > div { background:var(--surface) !important; border:1px solid var(--border) !important; border-radius:10px !important; color:var(--text) !important; }
.stTabs [data-baseweb="tab-list"] { background:var(--surface) !important; border-radius:10px !important; padding:3px !important; border:1px solid var(--border) !important; }
.stTabs [data-baseweb="tab"] { background:transparent !important; color:var(--muted) !important; border-radius:7px !important; }
.stTabs [aria-selected="true"] { background:var(--accent) !important; color:white !important; }

/* ── Review Section ── */
.review-section-title {
    font-family:'Syne',sans-serif; font-weight:800; font-size:1.3rem;
    text-align:center; margin-bottom:.3rem; color:var(--text);
}
.review-section-sub {
    text-align:center; color:var(--muted); font-size:13px; margin-bottom:1.2rem;
}
.reviews-grid { display:grid; grid-template-columns:1fr 1fr; gap:.9rem; margin:1rem 0 1.5rem; }
.review-card {
    background:var(--card); border:1px solid var(--border);
    border-radius:12px; padding:1rem 1.1rem;
    box-shadow:0 2px 10px #f5c51812; transition: border-color .2s;
}
.review-card:hover { border-color:#c9960c66; }
.review-stars { color:#f5c518; font-size:15px; margin-bottom:.4rem; letter-spacing:1px; }
.review-body { font-size:13px; color:var(--text); line-height:1.65; margin-bottom:.6rem; font-style:italic; }
.review-name { font-weight:700; font-size:12px; color:var(--accent); }
.review-role { font-size:11px; color:var(--muted); margin-top:1px; }
.review-form-card {
    background:var(--card); border:1px solid var(--border);
    border-radius:var(--radius); padding:1.4rem 1.5rem; margin-top:.5rem;
    box-shadow:0 2px 12px #f5c51815;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def load_email_db():
    if os.path.exists(EMAIL_DB_FILE):
        with open(EMAIL_DB_FILE) as f:
            return json.load(f)
    return {}

def email_used(email):
    return email.lower() in load_email_db()

def mark_email_used(email):
    db = load_email_db()
    db[email.lower()] = datetime.now().isoformat()
    with open(EMAIL_DB_FILE, "w") as f:
        json.dump(db, f)

def read_pdf(file_bytes):
    import io
    import pdfplumber
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text.strip()

def log_payment_interest(plan, email=""):
    try:
        requests.post(SHEET_SCRIPT_URL, json={
            "type": "payment_interest",
            "plan": plan,
            "email": email,
            "timestamp": datetime.now().isoformat()
        }, timeout=5)
    except:
        pass  # Silently fail — never break the app for this

def log_user_action(action, email="", extra=None):
    """
    Sends detailed tracking to Google Sheets so you can see exactly
    what each user did — not just who logged in with their email.
    action values: "ai_processed", "resume_downloaded", "report_downloaded",
                   "payment_click", "mentorship_interest", "mock_interview_interest"
    """
    payload = {
        "type": "user_action",
        "action": action,
        "email": email,
        "timestamp": datetime.now().isoformat(),
    }
    if extra:
        payload.update(extra)
    try:
        requests.post(SHEET_SCRIPT_URL, json=payload, timeout=5)
    except:
        pass

def fetch_jd_from_url(url):
    try:
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script","style","nav","footer","header"]):
            tag.decompose()
        lines = [l.strip() for l in soup.get_text(separator="\n").splitlines() if l.strip()]
        return "\n".join(lines[:300])
    except:
        return None

def call_groq(resume_text, jd_text, market_mode="🇮🇳 India (Naukri / LinkedIn India)"):
    client = groq.Groq(api_key=GROQ_API_KEY)

    if "India" in market_mode:
        market_instructions = """
INDIAN JOB MARKET RULES (very important):
- Optimise keywords specifically for Naukri.com and LinkedIn India ATS ranking
- Use Indian resume conventions: include notice period if mentioned, CTC in LPA format, percentage-based education scores
- Add Indian recruiter search terms naturally: "immediate joiner", "open to relocation", relevant Indian tech stack terms
- Keep declaration section if present in original resume
- Bullet points should include measurable Indian industry-standard metrics
- Summary should mention notice period and location preference if available
- Skills should match exactly what Indian recruiters search for on Naukri
"""
    else:
        market_instructions = """
INTERNATIONAL JOB MARKET RULES (very important):
- Optimise keywords for global ATS systems: Workday, Greenhouse, Lever, Indeed
- Use international resume conventions: no photo, no DOB, no declaration, clean 1-page preferred
- Salary references in annual USD/GBP format if mentioned
- Use action verbs and quantified achievements suited for western hiring managers
- Skills and tools should match global industry-standard terminology
- Summary should be punchy, achievement-focused, and ATS-friendly for international roles
"""

    prompt = f"""
You are an expert ATS resume specialist and career coach.
Analyze the resume against the job description carefully.
Return ONLY a JSON object. No text before or after. No markdown. Just pure JSON.

{market_instructions}

{{
  "candidate_name": "full name from resume",
  "email": "email from resume or empty string",
  "phone": "phone from resume or empty string",
  "location": "city from resume or empty string",
  "linkedin": "linkedin url from resume or empty string",
  "match_score": 75,
  "ats_keywords_found": 12,
  "ats_keywords_missing": 5,
  "strong_points": ["point 1","point 2","point 3","point 4","point 5"],
  "missing_skills": ["skill 1","skill 2","skill 3","skill 4"],
  "improvement_tips": ["tip 1","tip 2","tip 3","tip 4"],
  "summary": "2 to 3 sentence professional summary tailored to the job description",
  "work_experience": [{{
    "title": "job title exactly as in resume",
    "company": "company name exactly as in resume",
    "dates": "dates exactly as in resume",
    "location": "location exactly as in resume",
    "bullets": ["bullet rewritten with JD keywords","bullet","bullet"]
  }}],
  "projects": [{{
    "name": "project name",
    "bullets": ["project description bullet"]
  }}],
  "education": [{{
    "degree": "degree name exactly as in resume",
    "institution": "institution name exactly as in resume",
    "year": "year exactly as in resume",
    "cgpa": "cgpa or empty string"
  }}],
  "skills_technical": ["skill1","skill2","skill3"],
  "skills_tools": ["tool1","tool2","tool3"],
  "achievements": ["achievement 1","achievement 2","achievement 3"],
  "certifications": ["certification 1","certification 2"],
  "score_explanation": "2-3 sentences explaining exactly why the ATS score improved — mention specific keywords added, sections strengthened, and what made the biggest difference",
  "job_title_suggestions": ["Job Title 1","Job Title 2","Job Title 3","Job Title 4","Job Title 5"]
}}

STRICT RULES:
- Never fabricate any experience skills or education
- Only use information already present in the resume
- Rewrite bullet points using keywords from the job description
- Keep all dates company names and institutions exactly as original
- match_score must be a number not a string
- Return ONLY pure JSON nothing else

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}
"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5,
        max_tokens=3500
    )
    return response.choices[0].message.content

def parse_json(text):
    text = text.strip()
    if "```" in text:
        lines = [l for l in text.split("\n") if not l.strip().startswith("```")]
        text = "\n".join(lines)
    start = text.find("{")
    end = text.rfind("}") + 1
    if start != -1 and end > start:
        text = text[start:end]
    return json.loads(text)

def simple_ats_score(resume_text, jd_text):
    jd_words  = set(re.findall(r'\b[a-zA-Z]{4,}\b', jd_text.lower()))
    res_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', resume_text.lower()))
    common    = jd_words & res_words
    score     = int((len(common) / max(len(jd_words), 1)) * 100)
    return max(min(score, 75), 5)  # Cap before-score at 75 max

# ─────────────────────────────────────────────
# DOCUMENT HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   kwargs.get(side, 'nil'))
        tag.set(qn('w:sz'),    '0')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'auto')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_para(doc_or_cell, text, bold=False, size=10, color=None,
             italic=False, align=None, space_before=0, space_after=60):
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
        p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def section_heading(doc, text, color="2E75B6"):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def bullet_para(doc, text, size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.lstrip('•-– '))
    run.font.size = Pt(size)
    return p

# ─────────────────────────────────────────────
# BUILD RESUME DOCX
# ─────────────────────────────────────────────

def build_resume(data, watermark=False):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── Watermark header ──
    if watermark:
        hdr = doc.sections[0].header
        hp  = hdr.paragraphs[0]
        hp.clear()
        run = hp.add_run("⚡ FREE VERSION — UPGRADE AT RESUMEREFLECT.STREAMLIT.APP FOR CLEAN COPY")
        run.font.size  = Pt(7)
        run.font.color.rgb = RGBColor(0x9B, 0x59, 0xB6)
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Name banner ──
    name_tbl = doc.add_table(rows=1, cols=1)
    name_tbl.style = 'Table Grid'
    cell = name_tbl.cell(0, 0)
    set_cell_bg(cell, '1A1A2E')
    set_cell_borders(cell, top='nil', bottom='nil', left='nil', right='nil')
    cell.paragraphs[0].clear()

    np = cell.paragraphs[0]
    nr = np.add_run((data.get('candidate_name') or 'Candidate').upper())
    nr.bold = True
    nr.font.size = Pt(18)
    nr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    np.paragraph_format.space_before = Pt(8)
    np.paragraph_format.space_after  = Pt(2)

    contact_parts = [x for x in [
        data.get('email'), data.get('phone'),
        data.get('location'), data.get('linkedin')
    ] if x]
    cp = cell.add_paragraph("   |   ".join(contact_parts))
    cp.runs[0].font.size = Pt(8)
    cp.runs[0].font.color.rgb = RGBColor(0xA0, 0xC4, 0xFF)
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(8)

    doc.add_paragraph()

    # ── Summary ──
    if data.get('summary'):
        section_heading(doc, 'Professional Summary')
        p = doc.add_paragraph(data['summary'])
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(4)

    # ── Work Experience ──
    jobs = data.get('work_experience', [])
    if jobs:
        section_heading(doc, 'Work Experience')
        for job in jobs:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            r1 = p.add_run(job.get('title', ''))
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(f"  ·  {job.get('company', '')}")
            r2.bold = True
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            r3 = p.add_run(f"  ·  {job.get('dates', '')}  ·  {job.get('location', '')}")
            r3.italic = True
            r3.font.size = Pt(8.5)
            r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            for b in job.get('bullets', [])[:3]:
                bullet_para(doc, b)

    # ── Projects ──
    projects = data.get('projects', [])
    if projects:
        section_heading(doc, 'Projects')
        for proj in projects[:2]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(proj.get('name', ''))
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            for b in proj.get('bullets', [])[:2]:
                bullet_para(doc, b)

    # ── Skills ──
    tech  = data.get('skills_technical', [])
    tools = data.get('skills_tools', [])
    if tech or tools:
        section_heading(doc, 'Skills')
        skills_tbl = doc.add_table(rows=1, cols=2)
        skills_tbl.style = 'Table Grid'
        w = skills_tbl.columns[0].width

        lc = skills_tbl.cell(0, 0)
        rc = skills_tbl.cell(0, 1)
        set_cell_bg(lc, 'F5F7FA')
        set_cell_bg(rc, 'F5F7FA')
        set_cell_borders(lc, top='nil', bottom='nil', left='nil', right='nil')
        set_cell_borders(rc, top='nil', bottom='nil', left='nil', right='nil')

        lc.paragraphs[0].clear()
        lp = lc.paragraphs[0]
        lr = lp.add_run('Technical Skills')
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        lp2 = lc.add_paragraph("  •  ".join(tech))
        lp2.runs[0].font.size = Pt(9)

        rc.paragraphs[0].clear()
        rp = rc.paragraphs[0]
        rr = rp.add_run('Tools & Technologies')
        rr.bold = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        rp2 = rc.add_paragraph("  •  ".join(tools))
        rp2.runs[0].font.size = Pt(9)

    # ── Education ──
    edu = data.get('education', [])
    if edu:
        section_heading(doc, 'Education')
        for e in edu:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(2)
            r1 = p.add_run(e.get('degree', ''))
            r1.bold = True
            r1.font.size = Pt(10)
            extra = f"  ·  {e.get('institution', '')}  ·  {e.get('year', '')}"
            if e.get('cgpa'):
                extra += f"  ·  CGPA: {e['cgpa']}"
            r2 = p.add_run(extra)
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Achievements ──
    ach = data.get('achievements', [])
    if ach:
        section_heading(doc, 'Achievements')
        for a in ach[:3]:
            bullet_para(doc, a)

    # ── Certifications ──
    certs = data.get('certifications', [])
    if certs:
        section_heading(doc, 'Certifications')
        for c in certs[:3]:
            bullet_para(doc, c)

    # ── Footer ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '2')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'D0D8E8')
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run("Tailored with ResumeReflect  •  AI Resume Tool")
    run.italic = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# BUILD ANALYSIS DOCX
# ─────────────────────────────────────────────

def build_analysis(data):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    score = max(5, int(data.get('match_score') or 0))

    if score >= 70:
        sc, sl = '00A651', 'STRONG MATCH — Ready to apply!'
    elif score >= 40:
        sc, sl = 'E67E22', 'MODERATE MATCH — Few improvements needed'
    else:
        sc, sl = 'C0392B', 'KEEP BUILDING — Focus on missing skills first'

    # ── Banner ──
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1B4F72')
    set_cell_borders(cell, top='nil', bottom='nil', left='nil', right='nil')

    cell.paragraphs[0].clear()
    bp = cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br1 = bp.add_run("CVNIXO")
    br1.bold = True
    br1.font.size = Pt(18)
    br1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    br2 = bp.add_run("  ·  ATS RESUME ANALYSIS REPORT")
    br2.font.size = Pt(11)
    br2.font.color.rgb = RGBColor(0x90, 0xC4, 0xF0)
    bp.paragraph_format.space_before = Pt(8)
    bp.paragraph_format.space_after  = Pt(2)

    cp = cell.add_paragraph(f"Candidate: {data.get('candidate_name', '')}")
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(9)
    cp.runs[0].font.color.rgb = RGBColor(0xC8, 0xDF, 0xF5)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ── Score card ──
    stbl = doc.add_table(rows=1, cols=2)
    stbl.style = 'Table Grid'

    lc = stbl.cell(0, 0)
    rc = stbl.cell(0, 1)
    set_cell_bg(lc, 'F5F7FA')
    set_cell_bg(rc, 'FFFFFF')
    set_cell_borders(lc, top='nil', bottom='nil', left='nil', right='nil')
    set_cell_borders(rc, top='nil', bottom='nil', left='nil', right='nil')

    lc.paragraphs[0].clear()
    sp = lc.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"{score}%")
    sr.bold = True
    sr.font.size = Pt(36)
    sr.font.color.rgb = RGBColor(*bytes.fromhex(sc))
    sp.paragraph_format.space_before = Pt(10)
    sp.paragraph_format.space_after  = Pt(2)

    lp2 = lc.add_paragraph("ATS MATCH SCORE")
    lp2.runs[0].bold = True
    lp2.runs[0].font.size = Pt(7.5)
    lp2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lp3 = lc.add_paragraph(sl)
    lp3.runs[0].bold = True
    lp3.runs[0].font.size = Pt(8.5)
    lp3.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(sc))
    lp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp3.paragraph_format.space_after = Pt(10)

    rc.paragraphs[0].clear()
    rp = rc.paragraphs[0]
    rp.paragraph_format.space_before = Pt(10)
    rr1 = rp.add_run(f"Keywords Found: {data.get('ats_keywords_found', '?')}     ")
    rr1.bold = True
    rr1.font.size = Pt(9)
    rr1.font.color.rgb = RGBColor(0x00, 0xA6, 0x51)
    rr2 = rp.add_run(f"Keywords Missing: {data.get('ats_keywords_missing', '?')}")
    rr2.bold = True
    rr2.font.size = Pt(9)
    rr2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph()

    # ── Strong vs Missing ──
    section_heading(doc, 'Detailed Analysis', '1B4F72')
    atbl = doc.add_table(rows=2, cols=2)
    atbl.style = 'Table Grid'

    hdr_l = atbl.cell(0, 0)
    hdr_r = atbl.cell(0, 1)
    set_cell_bg(hdr_l, 'E8F5E9')
    set_cell_bg(hdr_r, 'FDEDEC')
    set_cell_borders(hdr_l, top='nil', bottom='nil', left='nil', right='nil')
    set_cell_borders(hdr_r, top='nil', bottom='nil', left='nil', right='nil')

    hdr_l.paragraphs[0].clear()
    hl = hdr_l.paragraphs[0]
    hlr = hl.add_run("✓  WHAT MATCHES WELL")
    hlr.bold = True
    hlr.font.size = Pt(9.5)
    hlr.font.color.rgb = RGBColor(0x00, 0xA6, 0x51)
    hl.paragraph_format.space_before = Pt(6)
    hl.paragraph_format.space_after = Pt(4)

    hdr_r.paragraphs[0].clear()
    hr = hdr_r.paragraphs[0]
    hrr = hr.add_run("✗  NEEDS IMPROVEMENT")
    hrr.bold = True
    hrr.font.size = Pt(9.5)
    hrr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    hr.paragraph_format.space_before = Pt(6)
    hr.paragraph_format.space_after = Pt(4)

    body_l = atbl.cell(1, 0)
    body_r = atbl.cell(1, 1)
    set_cell_bg(body_l, 'E8F5E9')
    set_cell_bg(body_r, 'FDEDEC')
    set_cell_borders(body_l, top='nil', bottom='nil', left='nil', right='nil')
    set_cell_borders(body_r, top='nil', bottom='nil', left='nil', right='nil')

    body_l.paragraphs[0].clear()
    for pt in data.get('strong_points', []):
        p = body_l.add_paragraph(f"• {pt.lstrip('•- ')}")
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)

    body_r.paragraphs[0].clear()
    for ms in data.get('missing_skills', []):
        p = body_r.add_paragraph(f"• {ms.lstrip('•- ')}")
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # ── Improvement Roadmap ──
    section_heading(doc, 'Improvement Roadmap', 'E67E22')
    for i, tip in enumerate(data.get('improvement_tips', []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{i}.  ")
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        r2 = p.add_run(tip)
        r2.font.size = Pt(9.5)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph("Generated by ResumeReflect  •  AI Resume Tailoring Tool")
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(7.5)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
defaults = {
    "step": 1, "resume_text": None, "jd_text": None,
    "ai_data": None, "ats_before": None, "ats_after": None,
    "plan": None, "email": None, "market_mode": "🇮🇳 India (Naukri / LinkedIn India)",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ─────────────────────────────────────────────
# HERO
# ─────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <div class="badge">⚡ AI Powered · ATS Optimized</div>
    <h1>Get Past the ATS.<br>Land the Interview.</h1>
    <p>ResumeReflect tailors your resume to any job in 30 seconds — precise, natural, never fabricated.</p>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# STEP 1 — RESUME
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step-num">1</span> Upload Your Resume (PDF)</div>', unsafe_allow_html=True)
resume_file = st.file_uploader("Resume", type=["pdf"], label_visibility="collapsed")
st.markdown('</div>', unsafe_allow_html=True)

if resume_file:
    file_bytes = resume_file.read()
    st.session_state.resume_text = read_pdf(file_bytes)
    st.success(f"✓ Resume loaded — {len(st.session_state.resume_text.split())} words")


# ─────────────────────────────────────────────
# MARKET SELECTOR
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title">🌍 Which job market are you targeting?</div>', unsafe_allow_html=True)
market_mode = st.radio("Market", [
    "🇮🇳 India (Naukri / LinkedIn India)",
    "🌍 International (LinkedIn / Indeed / Workday)"
], label_visibility="collapsed", horizontal=True)
st.session_state.market_mode = market_mode
st.markdown('</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# STEP 2 — JD
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step-num">2</span> Add the Job Description</div>', unsafe_allow_html=True)
jd_method = st.radio("JD method",
    ["Paste job link (LinkedIn / Naukri / Indeed / any)", "Paste JD text directly"],
    label_visibility="collapsed", horizontal=True)

if "link" in jd_method.lower():
    job_url = st.text_input("Job URL", placeholder="https://linkedin.com/jobs/view/...", label_visibility="collapsed")
    if job_url and st.button("🔗 Fetch Job Description", key="fetch"):
        with st.spinner("Fetching job details..."):
            fetched = fetch_jd_from_url(job_url)
        if fetched:
            st.session_state.jd_text = fetched
            st.success("✓ Job description fetched!")
        else:
            st.error("Could not fetch. Please paste the JD text directly.")
else:
    jd_raw = st.text_area("Paste full job description", height=180,
        placeholder="Paste the complete job description here...", label_visibility="collapsed")
    if jd_raw:
        st.session_state.jd_text = jd_raw
st.markdown('</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# STEP 3 — EMAIL + GENERATE
# ─────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step-num">3</span> Your Email — Get Free Tailored Resume</div>', unsafe_allow_html=True)
email_input = st.text_input("Email", placeholder="you@email.com", label_visibility="collapsed")
go = st.button("⚡ Tailor My Resume — Free")
st.markdown('</div>', unsafe_allow_html=True)

if go:
    if not st.session_state.resume_text:
        st.error("Please upload your resume first.")
    elif not st.session_state.jd_text:
        st.error("Please add the job description.")
    elif not email_input or "@" not in email_input:
        st.error("Please enter a valid email.")
    elif email_used(email_input):
        st.warning("This email has already used the free tier. Upgrade below to continue.")
        st.session_state.step = 2
    else:
        st.session_state.ats_before = simple_ats_score(st.session_state.resume_text, st.session_state.jd_text)

        with st.spinner("AI is analysing and tailoring your resume..."):
            raw = call_groq(st.session_state.resume_text, st.session_state.jd_text, st.session_state.market_mode)

        try:
            data = parse_json(raw)
        except Exception as e:
            st.error(f"Parsing error — please try again.")
            st.stop()

        st.session_state.ai_data   = data
        # Safely parse AI score — clamp to 95 max so bar never overflows 100%
        try:
            ai_score = int(float(str(data.get('match_score', 75))))
        except:
            ai_score = 75
        ai_score = max(1, min(ai_score, 95))  # hard clamp: 1–95

        before   = st.session_state.ats_before  # already 5–75 from simple_ats_score
        # After: at least before+5, never above 95
        st.session_state.ats_after = min(max(ai_score, before + 5), 95)
        st.session_state.email     = email_input
        mark_email_used(email_input)

        # ── Rich tracking: log exactly what AI found for this user ──
        log_user_action("ai_processed", email=email_input, extra={
            "candidate_name":  data.get("candidate_name", ""),
            "ats_before":      st.session_state.ats_before,
            "ats_after":       st.session_state.ats_after,
            "market_mode":     st.session_state.market_mode,
            "missing_skills":  ", ".join(data.get("missing_skills", [])[:4]),
            "match_score_raw": data.get("match_score", ""),
        })

        st.session_state.step = 2
        st.rerun()


# ─────────────────────────────────────────────
# RESULTS
# ─────────────────────────────────────────────
if st.session_state.step >= 2 and st.session_state.ai_data:

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    before = min(int(st.session_state.ats_before or 0), 100)
    after  = min(int(st.session_state.ats_after  or 0), 100)

    st.markdown(f"""
    <div class="score-wrap">
        <div class="score-box">
            <div class="score-val score-before">{before}%</div>
            <div class="bar-wrap"><div class="bar" style="width:{before}%"></div></div>
            <div class="score-label">Before Tailoring</div>
        </div>
        <div class="score-box">
            <div class="score-val score-after">{after}%</div>
            <div class="bar-wrap"><div class="bar" style="width:{after}%"></div></div>
            <div class="score-label">After Tailoring · ATS Match</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    data = st.session_state.ai_data

    # ── Score Explanation ──
    score_explanation = data.get('score_explanation', '')
    if score_explanation:
        improvement = after - before
        st.markdown(f"""
    <div class="card" style="border-color:#00e5a033;margin-top:.5rem">
        <div class="card-title">📈 Why Your Score Improved by {improvement}%</div>
        <div style="color:var(--text);font-size:14px;line-height:1.7">{score_explanation}</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Job Title Suggestions ──
    job_titles = data.get('job_title_suggestions', [])
    if job_titles:
        market = st.session_state.get('market_mode', '')
        platform = "Naukri / LinkedIn" if "India" in market else "LinkedIn / Indeed"
        chips = "".join([f'<span class="chip">🔍 {t}</span>' for t in job_titles])
        st.markdown(f"""
    <div class="card" style="border-color:#6c63ff33;margin-top:.5rem">
        <div class="card-title">💼 Job Titles to Search on {platform}</div>
        <div style="margin-top:.4rem">{chips}</div>
        <div style="color:var(--muted);font-size:11px;margin-top:.8rem">Based on your resume — search these exact titles for best results</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="divider"></div>
    <div class="divider"></div>
    """, unsafe_allow_html=True)

    # ── Mock Interview Card ──
    st.markdown("""
    <div class="card" style="border-color:#f59e0b44;margin-top:.5rem">
        <div style="display:flex;align-items:center;gap:.6rem;margin-bottom:.5rem">
            <span style="font-size:1.5rem">🎤</span>
            <div class="card-title" style="margin:0">Mock Interview with Real Human Feedback</div>
        </div>
        <div style="color:var(--muted);font-size:13px;line-height:1.7;margin-bottom:.8rem">
            Resume ready? Now prepare for the interview.<br>
            30-min 1-on-1 call · Real questions based on your resume & JD · Human feedback on communication, confidence & body language · AI cannot do this.
        </div>
        <div style="display:flex;gap:1rem;flex-wrap:wrap;margin-bottom:.6rem">
            <span style="background:#f59e0b22;color:#f59e0b;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">$19 / session</span>
            <span style="background:#00e5a022;color:#00e5a0;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">30 minutes</span>
            <span style="background:#6c63ff22;color:#a78bfa;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">Google Meet</span>
            <span style="background:#ff6b6b22;color:#ff6b6b;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">Launching Soon</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("🔔 Notify Me — Mock Interview", key="mock_interview_notify", use_container_width=True):
        log_payment_interest("Mock Interview $19", st.session_state.email or "")
        st.success("✅ You are on the list! We will notify you as soon as mock interviews go live.")

    # ── 6-Month Mentorship Card ──
    st.markdown("""
    <div class="card" style="border-color:#00e5a044;margin-top:1rem">
        <div style="display:flex;align-items:center;gap:.6rem;margin-bottom:.5rem">
            <span style="font-size:1.5rem">🚀</span>
            <div class="card-title" style="margin:0">6-Month Career Mentorship — Personal 1-on-1</div>
        </div>
        <div style="color:var(--muted);font-size:13px;line-height:1.7;margin-bottom:.8rem">
            Not just a tool — a real mentor by your side for 6 months.<br>
            Monthly strategy calls · Resume updates for every job · LinkedIn profile rewrite · Interview prep · Job search guidance · You get a dedicated career partner, not AI.
        </div>
        <div style="display:flex;gap:1rem;flex-wrap:wrap;margin-bottom:.6rem">
            <span style="background:#00e5a022;color:#00e5a0;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">$299 / 6 months</span>
            <span style="background:#f59e0b22;color:#f59e0b;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">1-on-1 Human Mentor</span>
            <span style="background:#6c63ff22;color:#a78bfa;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">WhatsApp + Google Meet</span>
            <span style="background:#ff6b6b22;color:#ff6b6b;padding:.2rem .7rem;border-radius:20px;font-size:11px;font-weight:600">Limited Spots</span>
        </div>
        <div style="color:var(--muted);font-size:11px;margin-top:.4rem">
            ✅ Resume tailored for every application &nbsp;|&nbsp; ✅ Mock interviews included &nbsp;|&nbsp; ✅ Real human support
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("🔔 Notify Me — 6-Month Mentorship", key="mentorship_notify", use_container_width=True):
        log_payment_interest("6-Month Mentorship $299", st.session_state.email or "")
        st.success("✅ Noted! We will contact you personally to discuss the mentorship program.")

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="card" style="border-color:#6c63ff33">
        <div class="card-title">✅ Your Documents Are Ready — Download Below</div>
        <div style="color:var(--muted);font-size:13px">
            Free version has a small watermark. Upgrade for clean copy + cover letter + interview kit.
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Generate docs
    resume_bytes   = build_resume(data, watermark=True)
    analysis_bytes = build_analysis(data)

    col1, col2 = st.columns(2)
    with col1:
        if st.download_button(
            "⬇ Download Tailored Resume (Free)",
            data=resume_bytes,
            file_name="resumereflect_resume_free.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            log_user_action("resume_downloaded", email=st.session_state.get("email",""))
    with col2:
        if st.download_button(
            "⬇ Download ATS Report (Free)",
            data=analysis_bytes,
            file_name="resumereflect_ats_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            log_user_action("report_downloaded", email=st.session_state.get("email",""))

    # ── PRICING ──
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align:center;margin-bottom:1.2rem">
        <div style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800">Unlock Full Power</div>
        <div style="color:var(--muted);font-size:13px;margin-top:.3rem">One payment. Instant unlock. No subscriptions.</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="pricing-grid">
        <div class="plan">
            <div class="plan-name">Basic</div>
            <div class="plan-price">$1</div>
            <div class="plan-period">one-time</div>
            <ul class="plan-feat">
                <li>Clean resume (no watermark)</li>
                <li>ATS Score Report</li>
                <li>DOCX format</li>
            </ul>
        </div>
        <div class="plan hot">
            <div class="hot-badge">POPULAR</div>
            <div class="plan-name">Pro</div>
            <div class="plan-price">$19</div>
            <div class="plan-period">one-time</div>
            <ul class="plan-feat">
                <li>Everything in Basic</li>
                <li>Cover Letter</li>
                <li>Interview Prep Kit</li>
            </ul>
        </div>
        <div class="plan">
            <div class="plan-name">Monthly</div>
            <div class="plan-price">$29</div>
            <div class="plan-period">per month · unlimited</div>
            <ul class="plan-feat">
                <li>Everything in Pro</li>
                <li>LinkedIn Profile Rewrite</li>
                <li>Unlimited resumes</li>
            </ul>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Coming Soon Pay Buttons ──
    col_b, col_p, col_m = st.columns(3)
    with col_b:
        if st.button("Pay $1 →", key="pay_basic", use_container_width=True):
            log_payment_interest("Basic $1", st.session_state.email or "")
            st.success("🚀 Payments launching soon! We've noted your interest.")
    with col_p:
        if st.button("Pay $19 →", key="pay_pro", use_container_width=True):
            log_payment_interest("Pro $19", st.session_state.email or "")
            st.success("🚀 Payments launching soon! We've noted your interest.")
    with col_m:
        if st.button("Pay $29 →", key="pay_monthly", use_container_width=True):
            log_payment_interest("Monthly $29", st.session_state.email or "")
            st.success("🚀 Payments launching soon! We've noted your interest.")

    # ── UNLOCK ──
    st.markdown("""
    <div class="unlock-box">
        <div class="card-title">✅ Already Paid? Enter Transaction ID to Unlock</div>
        <div style="color:var(--muted);font-size:12px;margin-bottom:.8rem">
            After payment your UPI app shows a Transaction ID — paste it below. Instant unlock, no email needed.
        </div>
    </div>
    """, unsafe_allow_html=True)

    c1, c2 = st.columns([3, 1])
    with c1:
        txn = st.text_input("Transaction ID", placeholder="UPI ref number or Razorpay payment ID", label_visibility="collapsed")
    with c2:
        plan = st.selectbox("Plan", ["Basic $1", "Pro $19", "Monthly $29"], label_visibility="collapsed")

    if st.button("🔓 Unlock My Plan"):
        if txn and len(txn) > 6:
            st.session_state.plan = plan
            st.session_state.step = 3
            st.rerun()
        else:
            st.error("Please enter your transaction ID from your UPI app.")


# ─────────────────────────────────────────────
# UNLOCKED
# ─────────────────────────────────────────────
if st.session_state.step == 3 and st.session_state.ai_data:

    plan      = st.session_state.plan or ""
    is_pro    = "Pro" in plan or "Yearly" in plan
    is_yearly = "Monthly" in plan
    data      = st.session_state.ai_data

    st.success(f"🎉 {plan} unlocked! Download your files below.")

    # Clean resume
    clean_bytes = build_resume(data, watermark=False)
    st.download_button(
        "⬇ Download Clean Resume (No Watermark)",
        data=clean_bytes,
        file_name="resumereflect_tailored_resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if is_pro:
        client = groq.Groq(api_key=GROQ_API_KEY)
        tabs = st.tabs(["📝 Cover Letter", "🎯 Interview Kit",
                        "💼 LinkedIn" if is_yearly else "💼 LinkedIn (Yearly only)"])

        with tabs[0]:
            with st.spinner("Writing cover letter..."):
                cl = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role":"user","content":f"""
Write a compelling, personalized cover letter. 3 paragraphs. Professional but human tone. No generic openers like "I am writing to express".
Candidate summary: {data.get('summary','')}
Strong points: {', '.join(data.get('strong_points',[]))}
Role context: {st.session_state.jd_text[:800] if st.session_state.jd_text else ''}
"""}],
                    temperature=0.6, max_tokens=700
                )
            cover = cl.choices[0].message.content
            st.text_area("Cover Letter", cover, height=280)
            st.download_button("⬇ Download Cover Letter", data=cover, file_name="resumereflect_cover_letter.txt", mime="text/plain")

        with tabs[1]:
            with st.spinner("Building interview kit..."):
                ik = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role":"user","content":f"""
Create an interview prep kit with:
1. TOP 5 LIKELY INTERVIEW QUESTIONS with brief answer hints based on the resume
2. 3 SMART QUESTIONS TO ASK THE INTERVIEWER
3. KEY SKILLS TO HIGHLIGHT IN THE INTERVIEW

Strong points: {', '.join(data.get('strong_points',[]))}
Missing skills: {', '.join(data.get('missing_skills',[]))}
Role context: {st.session_state.jd_text[:800] if st.session_state.jd_text else ''}
"""}],
                    temperature=0.5, max_tokens=900
                )
            kit = ik.choices[0].message.content
            st.text_area("Interview Kit", kit, height=350)
            st.download_button("⬇ Download Interview Kit", data=kit, file_name="resumereflect_interview_kit.txt", mime="text/plain")

        with tabs[2]:
            if is_yearly:
                with st.spinner("Rewriting LinkedIn profile..."):
                    li = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role":"user","content":f"""
Rewrite this person's LinkedIn profile. Provide:
1. HEADLINE (120 chars max, keyword-rich, compelling)
2. ABOUT SECTION (first-person, engaging, 1500-2000 chars)
3. TOP 5 SKILLS TO ADD ON LINKEDIN

Summary: {data.get('summary','')}
Technical skills: {', '.join(data.get('skills_technical',[]))}
Experience: {', '.join([j.get('title','') + ' at ' + j.get('company','') for j in data.get('work_experience',[])])}
"""}],
                        temperature=0.6, max_tokens=900
                    )
                linkedin = li.choices[0].message.content
                st.text_area("LinkedIn Rewrite", linkedin, height=350)
                st.download_button("⬇ Download LinkedIn Rewrite", data=linkedin, file_name="resumereflect_linkedin.txt", mime="text/plain")
            else:
                st.info("LinkedIn Profile Rewrite is available in the Monthly plan ($29). Upgrade to unlock.")

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    if st.button("🔄 Tailor Another Resume"):
        for k, v in defaults.items():
            st.session_state[k] = v
        st.session_state.step = 1
        st.rerun()


# ─────────────────────────────────────────────
# REVIEW SECTION
# ─────────────────────────────────────────────
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

st.markdown("""
<div class="review-section-title">⭐ What Users Are Saying</div>
<div class="review-section-sub">Real people. Real results. Join hundreds of job seekers who landed interviews.</div>

<div class="reviews-grid">
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"My ATS score jumped from 34% to 78% after tailoring. Got a callback from TCS within 3 days. This tool is unreal."</div>
        <div class="review-name">Arjun M.</div>
        <div class="review-role">Software Engineer · Bangalore</div>
    </div>
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"I was applying for months with no response. After ResumeReflect, I got 2 interview calls in the first week. Highly recommend!"</div>
        <div class="review-name">Priya S.</div>
        <div class="review-role">Data Analyst · Mumbai</div>
    </div>
    <div class="review-card">
        <div class="review-stars">★★★★☆</div>
        <div class="review-body">"The ATS report showed me exactly what keywords I was missing. Tailored my resume in 2 minutes and landed a Wipro interview."</div>
        <div class="review-name">Rohit K.</div>
        <div class="review-role">Business Analyst · Hyderabad</div>
    </div>
    <div class="review-card">
        <div class="review-stars">★★★★★</div>
        <div class="review-body">"As a fresher I had no idea how ATS works. This tool explained everything and rewrote my resume perfectly for my target roles."</div>
        <div class="review-name">Sneha T.</div>
        <div class="review-role">Marketing Executive · Pune</div>
    </div>
</div>
""", unsafe_allow_html=True)

# ── Leave a review form ──
st.markdown("""
<div class="review-form-card">
    <div class="card-title">✍️ Share Your Experience</div>
    <div style="color:var(--muted);font-size:12px;margin-bottom:1rem">
        Got an interview after using ResumeReflect? Your story helps others — and takes 60 seconds!
    </div>
</div>
""", unsafe_allow_html=True)

r_col1, r_col2 = st.columns(2)
with r_col1:
    reviewer_name = st.text_input(
        "Your Name",
        placeholder="e.g. Rahul S. (or leave blank for Anonymous)",
        key="reviewer_name"
    )
with r_col2:
    reviewer_role = st.text_input(
        "Job Role You Applied For",
        placeholder="e.g. Software Engineer at Infosys",
        key="reviewer_role"
    )

review_text = st.text_area(
    "Your Experience",
    placeholder="How did ResumeReflect help you? Did you get an interview call? What was your ATS score before/after? Any feedback for us?",
    height=110,
    key="review_text"
)

review_rating = st.select_slider(
    "⭐ Your Rating",
    options=["1 Star", "2 Stars", "3 Stars", "4 Stars", "5 Stars"],
    value="5 Stars",
    key="review_rating"
)

if st.button("Submit My Review ✅", key="submit_review", use_container_width=True):
    if review_text and len(review_text.strip()) > 10:
        try:
            requests.post(SHEET_SCRIPT_URL, json={
                "type":      "review",
                "name":      reviewer_name.strip() if reviewer_name.strip() else "Anonymous",
                "role":      reviewer_role or "",
                "review":    review_text.strip(),
                "rating":    review_rating,
                "email":     st.session_state.get("email", ""),
                "timestamp": datetime.now().isoformat()
            }, timeout=5)
        except:
            pass
        st.success("🙏 Thank you! Your review has been submitted. It really helps the community.")
        st.balloons()
    else:
        st.error("Please write at least a sentence about your experience.")

st.markdown('<div style="margin-top:1.5rem"></div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown("""
<div style="text-align:center;padding:2.5rem 0 1rem;color:var(--muted);font-size:12px">
    <div style="font-family:'Syne',sans-serif;font-weight:800;font-size:1rem;
        background:linear-gradient(135deg,var(--accent),var(--accent2));
        -webkit-background-clip:text;-webkit-text-fill-color:transparent;
        background-clip:text;margin-bottom:.4rem">⚡ ResumeReflect</div>
    Built for serious job seekers · Made in India 🇮🇳
    <br><br>
    <div style="font-size:11px;margin-bottom:.3rem">🔒 Your resume is processed by AI and not stored on our servers.</div>
    <div style="font-size:11px">© 2026 ResumeReflect. All rights reserved.</div>
</div>
""", unsafe_allow_html=True)

